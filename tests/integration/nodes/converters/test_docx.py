from io import BytesIO

import pytest
from docx import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from dynamiq import Workflow
from dynamiq.components.converters.docx import DOCXConverter
from dynamiq.flows import Flow
from dynamiq.nodes.converters.docx import DOCXFileConverter, DOCXFileConverterInputSchema
from dynamiq.runnables import RunnableResult, RunnableStatus
from dynamiq.types import Document


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def test_workflow_with_docx_converter():
    content = "Hello, World!"

    doc = DocxDocument()
    doc.add_paragraph(content)

    docx_converter = DOCXFileConverter()
    wf_docx = Workflow(flow=Flow(nodes=[docx_converter]))
    file = BytesIO()
    doc.save(file)
    file.name = "mock.docx"
    input_data = {"files": [file]}

    response = wf_docx.run(input_data=input_data)
    document_id = response.output[next(iter(response.output))]["output"]["documents"][0]["id"]
    docx_converter_expected_result = RunnableResult(
        status=RunnableStatus.SUCCESS,
        input=dict(DOCXFileConverterInputSchema(**input_data)),
        output={
            "documents": [
                Document(id=document_id, content=content, metadata={"file_path": file.name, "source": file.name})
            ]
        },
    ).to_dict(skip_format_types={BytesIO, bytes})

    expected_output = {docx_converter.id: docx_converter_expected_result}
    assert response == RunnableResult(status=RunnableStatus.SUCCESS, input=input_data, output=expected_output)


def test_docx_converter_preserves_run_spacing_hyperlinks_and_table_cells():
    doc = DocxDocument()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Read ")
    add_hyperlink(paragraph, "the docs", "https://example.com/docs")
    paragraph.add_run(" now.")
    formatted_paragraph = doc.add_paragraph()
    formatted_run = formatted_paragraph.add_run(" important ")
    formatted_run.bold = True
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Feature"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "A | B"
    table.cell(1, 1).text = "Line 1\nLine 2"
    file = BytesIO()
    doc.save(file)
    file.name = "guide.docx"

    converter = DOCXFileConverter()
    converter.init_components()
    result = converter.run(input_data={"files": [file]})

    assert result.status == RunnableStatus.SUCCESS
    content = result.output["documents"][0].content
    assert "Read [the docs](https://example.com/docs) now." in content
    assert " **important** " in content
    assert "| A \\| B | Line 1<br>Line 2 |" in content


def test_docx_legacy_page_mode_reports_real_section_numbers():
    doc = DocxDocument()
    doc.add_paragraph("First section")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section")
    file = BytesIO()
    doc.save(file)
    file.name = "sections.docx"

    converter = DOCXFileConverter(document_creation_mode="one-doc-per-page")
    converter.init_components()
    result = converter.run(input_data={"files": [file]})

    assert result.status == RunnableStatus.SUCCESS
    documents = result.output["documents"]
    assert [document.content for document in documents] == ["First section", "Second section"]
    assert [document.metadata["section_number"] for document in documents] == [1, 2]
    assert all("page_number" not in document.metadata for document in documents)


def test_docx_section_mode_rejects_document_without_extractable_content():
    doc = DocxDocument()
    file = BytesIO()
    doc.save(file)
    file.name = "empty.docx"

    converter = DOCXConverter(document_creation_mode="one-doc-per-page")

    with pytest.raises(ValueError, match="contains no extractable content"):
        converter.run(files=[file])
