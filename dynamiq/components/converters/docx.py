import copy
import re
from io import BytesIO
from pathlib import Path
from typing import Any, Literal

from docx import Document as DocxDocument
from docx.table import Table
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from dynamiq.components.converters.base import BaseConverter
from dynamiq.components.converters.utils import build_source_metadata, get_filename_for_bytesio
from dynamiq.types import Document, DocumentCreationMode, DocumentType

InlineElement = tuple[str, dict | None, str | None]
MARKDOWN_FORMAT_MARKERS = (
    ("bold", "**"),
    ("italic", "*"),
    ("underline", "_"),
)


class DOCXConverter(BaseConverter):
    """
    A component for converting DOCX files to Documents using the python-docx library.

    Initializes the object with the configuration for converting documents using
    python-docx.

    Args:
        document_creation_mode (Literal["one-doc-per-file", "one-doc-per-page"], optional):
            Determines how to create Documents from the elements of the Word document. Options are:
            - `"one-doc-per-file"`: Creates one Document per file.
                All elements are concatenated into one text field.
            - `"one-doc-per-page"`: Legacy name for splitting on DOCX section breaks.
                DOCX does not contain stable rendered page boundaries.
            Defaults to `"one-doc-per-file"`.

    Usage example:
        ```python
        from dynamiq.components.converters.docx import DOCXConverter

        converter = DOCXConverter()
        documents = converter.run(paths=["a/file/path.docx", "a/directory/path"])["documents"]
        ```
    """

    document_creation_mode: Literal[DocumentCreationMode.ONE_DOC_PER_FILE, DocumentCreationMode.ONE_DOC_PER_PAGE] = (
        DocumentCreationMode.ONE_DOC_PER_FILE
    )

    xml_key: str = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
    xml_namespaces: dict[str, str] = {
        "w": "http://schemas.microsoft.com/office/word/2003/wordml",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    image_placeholder: str = "<!-- image -->"
    blip_tag: str = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"

    def _process_file(self, file: Path | BytesIO, metadata: dict[str, Any]) -> list[Any]:
        """
        Process a single Word document and create documents.

        Args:
            file (Union[Path, BytesIO]): The file to process.
            metadata (Dict[str, Any]): Metadata to attach to the documents.

        Returns:
            List[Any]: A list of created documents.

        Raises:
            TypeError: If the file argument is neither a Path nor a BytesIO object.
        """
        if isinstance(file, Path):
            with open(file, "rb") as upload_file:
                file_content = BytesIO(upload_file.read())
                file_path = upload_file.name
        elif isinstance(file, BytesIO):
            file_path = get_filename_for_bytesio(file)
            file_content = file
        else:
            raise TypeError("Expected a Path object or a BytesIO object.")

        file_content.seek(0)
        elements = DocxDocument(file_content)
        return self._create_documents(
            filepath=file_path,
            elements=elements,
            document_creation_mode=self.document_creation_mode,
            metadata=metadata,
        )

    def _create_documents(
        self,
        filepath: str,
        elements: DocxDocument,
        document_creation_mode: DocumentCreationMode,
        metadata: dict[str, Any],
        **kwargs,
    ) -> list[Document]:
        """
        Create Documents from the elements of the Word document.
        """
        docs = []
        if document_creation_mode == DocumentCreationMode.ONE_DOC_PER_FILE:
            text_content = []

            for element in elements.element.body:
                tag_name = element.tag.split("}", 1)[-1] if "}" in element.tag else element.tag

                # Check for inline images
                contains_blip = any(child.tag == self.blip_tag for child in element.iter())

                if element.tag.endswith("p"):
                    text = self._process_paragraph(element, elements)
                    if contains_blip:
                        text = f"{text}\n\n{self.image_placeholder}" if text.strip() else self.image_placeholder
                    if text.strip():
                        text_content.append(text)
                elif element.tag.endswith("tbl"):
                    text = self._process_table(element, elements)
                    if text.strip():
                        text_content.append(text)
                elif contains_blip:
                    text_content.append(self.image_placeholder)
                elif tag_name in ["sdt"]:
                    sdt_content = element.find(".//w:sdtContent", namespaces=self.xml_namespaces)
                    if sdt_content is not None:
                        paragraphs = sdt_content.findall(".//w:p", namespaces=self.xml_namespaces)
                        for p in paragraphs:
                            text = self._process_paragraph(p, elements)
                            if text.strip():
                                text_content.append(text)

            full_text = "\n\n".join(text_content)
            if not full_text.strip():
                raise ValueError(f"DOCX file '{filepath}' contains no extractable content.")

            metadata = build_source_metadata(metadata, filepath)
            docs = [Document(content=full_text, metadata=metadata)]

        elif document_creation_mode == DocumentCreationMode.ONE_DOC_PER_PAGE:
            sections = self._split_into_sections(elements)
            if not sections:
                raise ValueError(f"DOCX file '{filepath}' contains no extractable content.")
            metadata = build_source_metadata(metadata, filepath)

            for idx, section_content in enumerate(sections, start=1):
                section_metadata = copy.deepcopy(metadata)
                # Keep the legacy key for consumers of the existing one-doc-per-page API.
                section_metadata["page_number"] = idx
                section_metadata["section_number"] = idx
                section_metadata["document_type"] = DocumentType.SECTION.value
                docs.append(Document(content=section_content, metadata=section_metadata))

        return docs

    def _process_paragraph(self, paragraph_element, docx_obj) -> str:
        """
        Process a paragraph element with formatting and hyperlinks.
        """
        paragraph = Paragraph(paragraph_element, docx_obj)
        paragraph_elements = self._get_paragraph_elements(paragraph)
        is_list_item, list_marker, list_level = self._check_for_list_item(paragraph)
        is_header, header_level = self._check_for_header(paragraph)
        inline_text = self._render_paragraph_elements(paragraph_elements)

        if is_header:
            header_prefix = "#" * header_level
            return f"{header_prefix} {inline_text}"
        if is_list_item:
            indent = "  " * (list_level - 1) if list_level > 1 else ""
            return f"{indent}{list_marker} {inline_text}"
        return inline_text

    @classmethod
    def _render_paragraph_elements(cls, paragraph_elements: list[InlineElement]) -> str:
        return "".join(cls._render_inline_element(*element) for element in paragraph_elements if element[0])

    @staticmethod
    def _render_inline_element(text: str, format_info: dict | None, hyperlink: str | None) -> str:
        """Render one DOCX inline element without placing Markdown markers around whitespace."""
        if not text.strip():
            return text
        leading_space = text[: len(text) - len(text.lstrip())]
        trailing_space = text[len(text.rstrip()) :]
        rendered = text.strip()

        for style, marker in MARKDOWN_FORMAT_MARKERS:
            if rendered and format_info and format_info.get(style, False):
                rendered = f"{marker}{rendered}{marker}"

        if rendered and hyperlink:
            rendered = f"[{rendered}]({hyperlink})"
        return f"{leading_space}{rendered}{trailing_space}"

    def _process_table(self, table_element, docx_obj) -> str:
        """
        Process a table element with enhanced features like cell spanning and formatting.
        """
        table = Table(table_element, docx_obj)

        # Check for single-cell tables
        if len(table.rows) == 1 and len(table.columns) == 1:
            cell_text = table.rows[0].cells[0].text
            if cell_text.strip():
                return cell_text

        table_rows = []

        if table.rows:
            header_cells = []
            for cell in table.rows[0].cells:
                header_cells.append(self._escape_table_cell(cell.text))
            table_rows.append("| " + " | ".join(header_cells) + " |")
            table_rows.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

        cell_set = set()
        for row_idx, row in enumerate(table.rows):
            # Skip the header row
            if row_idx == 0:
                continue

            row_cells = []
            for cell in row.cells:
                if cell._tc in cell_set:
                    continue
                cell_set.add(cell._tc)

                cell_text = self._escape_table_cell(cell.text)
                row_cells.append(cell_text)

            if row_cells:
                table_rows.append("| " + " | ".join(row_cells) + " |")

        return "\n".join(table_rows)

    @staticmethod
    def _escape_table_cell(value: str) -> str:
        return value.strip().replace("|", "\\|").replace("\n", "<br>")

    def _get_paragraph_elements(self, paragraph):
        """
        Extract paragraph elements (with the formatting and hyperlinks).
        """
        if paragraph.text.strip() == "":
            return [("", None, None)]

        paragraph_elements = []

        for content in paragraph.iter_inner_content():
            if isinstance(content, Hyperlink):
                text = content.text
                hyperlink = content.address if hasattr(content, "address") else None
                format_info = self._get_format_from_run(content.runs[0] if content.runs else None)
                if text:
                    paragraph_elements.append((text, format_info, hyperlink))
            elif isinstance(content, Run):
                text = content.text
                format_info = self._get_format_from_run(content)
                if text:
                    paragraph_elements.append((text, format_info, None))

        if not paragraph_elements and paragraph.text.strip():
            paragraph_elements.append((paragraph.text.strip(), None, None))

        return paragraph_elements

    def _get_format_from_run(self, run):
        """
        Extract formatting information from a run.
        """
        if not run:
            return None

        format_info = {}
        if hasattr(run, "bold") and run.bold:
            format_info["bold"] = True
        if hasattr(run, "italic") and run.italic:
            format_info["italic"] = True
        if hasattr(run, "underline") and run.underline:
            format_info["underline"] = True

        return format_info if format_info else None

    def _check_for_list_item(self, paragraph):
        """
        Check if a paragraph is a list item and return relevant information.
        """
        numbering_properties = paragraph._element.find(".//w:numPr", namespaces=paragraph._element.nsmap)
        if numbering_properties is not None:
            numbering_id_elem = numbering_properties.find("w:numId", namespaces=paragraph._element.nsmap)
            indexing_level_elem = numbering_properties.find("w:ilvl", namespaces=paragraph._element.nsmap)

            numbering_id = numbering_id_elem.get(self.xml_key) if numbering_id_elem is not None else None
            indexing_level = indexing_level_elem.get(self.xml_key) if indexing_level_elem is not None else None

            if numbering_id and numbering_id != "0":
                level = int(indexing_level) + 1 if indexing_level else 1
                # Check if the paragraph is a numbered list item or a bullet list item
                is_numbered = "Number" in paragraph.style.name if paragraph.style and paragraph.style.name else False
                marker = f"{level}." if is_numbered else "•"
                return True, marker, level

        return False, "", 0

    def _check_for_header(self, paragraph):
        """
        Check if a paragraph is a header and return the level.
        """
        if not paragraph.style:
            return False, 0

        style_id = paragraph.style.name.lower() if paragraph.style.name else ""

        heading_string = "heading"
        heading_pattern = r"heading\s*(\d+)"

        if heading_string in style_id:
            match = re.search(heading_pattern, style_id)
            if match:
                level = int(match.group(1))
                return True, level

        return False, 0

    def _split_into_sections(self, doc: DocxDocument) -> list[str]:
        """
        Split the document into sections based on section breaks.
        """
        sections = []
        current_section = []

        for element in doc.element.body:
            if element.tag.endswith("sectPr"):
                if current_section:
                    sections.append("\n\n".join(current_section))
                    current_section = []
            elif element.tag.endswith("p"):
                text = self._process_paragraph(element, doc)
                contains_blip = any(child.tag == self.blip_tag for child in element.iter())
                if contains_blip:
                    text = f"{text}\n\n{self.image_placeholder}" if text.strip() else self.image_placeholder
                if text.strip():
                    current_section.append(text)
                section_break = element.find(".//w:sectPr", namespaces=element.nsmap)
                if section_break is not None and current_section:
                    sections.append("\n\n".join(current_section))
                    current_section = []
            elif element.tag.endswith("tbl"):
                text = self._process_table(element, doc)
                if text.strip():
                    current_section.append(text)

        if current_section:
            sections.append("\n\n".join(current_section))
        return sections
